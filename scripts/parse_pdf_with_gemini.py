#!/usr/bin/env python3
"""
Script to parse quiz pack PDF files using Gemini API and convert to JSON format
Usage: python parse_pdf_with_gemini.py <pdf_file_path>
"""

import os
import sys
import json
import time
from google import genai
from google.genai import types
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    print("Error: GEMINI_API_KEY environment variable not set")
    print("Please set it in .env file or with: export GEMINI_API_KEY='your-api-key'")
    sys.exit(1)

client = genai.Client(api_key=GEMINI_API_KEY)

SYSTEM_PROMPT = """Ты эксперт по парсингу PDF файлов с вопросами для игры "Своя игра" в формат JSON.

═══════════════════════════════════════════════════════════════
КРИТИЧЕСКИ ВАЖНО! ЛОГИКА ПОЛЯ "FORM":
═══════════════════════════════════════════════════════════════

⚠️ ПОЛЕ "form" ≠ ОТВЕТ! ⚠️

"form" - это маркер/указание ЧТО именно нужно назвать, извлекается ИЗ САМОГО ВОПРОСА.

═══════════════════════════════════════════════════════════════
АЛГОРИТМ ИЗВЛЕЧЕНИЯ "FORM":
═══════════════════════════════════════════════════════════════

1️⃣ ИЩЕМ ЗАГЛАВНЫЕ СЛОВА/МЕСТОИМЕНИЯ В ВОПРОСЕ:
   - "В ЧЕСТЬ НЕГО назвали ящерицу" → form: "в честь него"
   - "ЭТОГО ГОЛЛАНДЦА упоминают" → form: "голландец" / "этого голландца"
   - "ОН шутил, что хотел" → form: "он"
   - "ЕГО хотели создать" → form: "его"
   - "ОНИ подвергались гонениям" → form: "они"
   - "С НЕЙ связано много легенд" → form: "она"
   - "ИМИ оказываются девушки" → form: "ими"
   - "ТАКИЕ ОНИ используются в игре" → form: "такие они"
   - "ЭТОЙ СТОЛИЦЕЙ является" → form: "столица" / "эта столица"
   - "ЭТОТ ГОРОД стал побратимом" → form: "этот город" / "город"
   - "В ЭТОМ ГОДУ погибла Ида" → form: "год"
   - "ЭТОГО ПРОДЮСЕРА работой было" → form: "продюсер"

2️⃣ ДОБАВЛЯЕМ ОГРАНИЧЕНИЯ ИЗ НАЧАЛА ВОПРОСА:
   - "В ответе одно слово. ОНА называется" → form: "одним словом, она"
   - "В ответе два слова. ОН известен" → form: "два слова, он"
   - "В ответе три слова. ЕГО называют" → form: "тремя словами, его"
   - "В ответе трёхсложное слово. ОНО" → form: "трёхсложное слово, оно"
   - "В ответе два слова на одну букву" → form: "два слова на одну букву, им"
   - "В ответе два слова на парные согласные" → form: "два слова на парные согласные"
   - "Зачет абсолютно точный ответ" → добавь это в form

3️⃣ ИСПОЛЬЗУЕМ ТИП ОБЪЕКТА ИЗ ВОПРОСА (если указан явно):
   - "ЭТОТ ФРАНЦУЗ не обладал" → form: "француз"
   - "ЭТА СТОЛИЦА находится" → form: "столица"
   - "ЭТОТ ГОРОД США стал" → form: "этот город"
   - "ЭТИМ АНГЛИЙСКИМ ВЫРАЖЕНИЕМ называют" → form: "английское выражение"
   - "ЭТОМУ ПРОИЗВЕДЕНИЮ посвящено" → form: "произведение"

4️⃣ ЕСЛИ В PDF ЕСТЬ СТРОКА "Зачет:":
   - Извлеки альтернативные варианты в отдельное поле "accept": ["вариант1", "вариант2"]
   - НО form всё равно должна быть местоимением/маркером из вопроса!

═══════════════════════════════════════════════════════════════
РЕАЛЬНЫЕ ПРИМЕРЫ ИЗ ПАКЕТОВ:
═══════════════════════════════════════════════════════════════

✅ ПРАВИЛЬНО:
Вопрос: "В декабре 2012 года В ЧЕСТЬ НЕГО назвали ископаемую ящерицу"
→ form: "в честь него"
→ answer: "Барак Обама"

✅ ПРАВИЛЬНО:
Вопрос: "Эстонец упоминается наравне с фамилией ЭТОГО ГОЛЛАНДЦА"
→ form: "голландец"
→ answer: "Ян Хендрик Оорт"

✅ ПРАВИЛЬНО:
Вопрос: "В ответе одно слово. «Третьей ЕЙ» называют загрязнение"
→ form: "одним словом, она"
→ answer: "рука"

✅ ПРАВИЛЬНО:
Вопрос: "ОН был профессором кафедры биохимии"
→ form: "он"
→ answer: "Александр Опарин"

✅ ПРАВИЛЬНО:
Вопрос: "ЭТА СТОЛИЦА находится в 25 км от экватора"
→ form: "столица"
→ answer: "Кито"

✅ ПРАВИЛЬНО:
Вопрос: "Девушки оказываются ИМИ"
→ form: "ими"
→ answer: "вампиры"

❌ НЕПРАВИЛЬНО:
→ form: "Барак Обама" ❌ (это ответ, не форма!)
→ form: "ответ" ❌ (бессмысленно)

═══════════════════════════════════════════════════════════════
ФОРМАТ ВЫХОДНОГО JSON:
═══════════════════════════════════════════════════════════════

{
  "info": "полная информация о пакете: название, авторы, редакторы, благодарности, мораторий",
  "package_name": "название пакета",
  "themes": [
    {
      "name": "Название темы (Автор: Имя Фамилия)",
      "theme_comment": "комментарий к теме, если есть в документе (НЕ выдумывать!)",
      "questions": [
        {
          "cost": 10,
          "question": "текст вопроса с ЗАГЛАВНЫМИ местоимениями",
          "form": "местоимение/маркер из вопроса (ОН/ОНА/ЕГО/ИХ/город/столица и т.д.)",
          "answer": "сам ответ",
          "source": "источник ответа (если указан)",
          "comment": "комментарий к ответу",
          "accept": ["альтернатива1", "альтернатива2"]  // опционально, если есть "Зачет:"
        }
      ]
    }
  ]
}

═══════════════════════════════════════════════════════════════
ОБЩИЕ ПРАВИЛА ПАРСИНГА:
═══════════════════════════════════════════════════════════════

1. Структура: Бои/Раунды → Темы → Вопросы (10, 20, 30, 40, 50 очков)
2. Info извлекается из начала файла (авторы, редакторы, благодарности, мораторий)
3. Сохраняй все ударения, специальные символы, форматирование
4. Поддерживай белорусский язык ("Адказ:", "Каментар:")
5. Если в вопросе несколько заглавных местоимений, выбирай основное (обычно последнее перед вопросительным знаком)
6. ВСЕГДА ищи и добавляй "source" (источник ответа) если он указан в документе:
   - Источник обычно после ответа в скобках: "(источник: ...)", "Источник: ...", "[...]"
   - Если источника нет, можно опустить поле
7. КРИТИЧНО: В JSON строках ОБЯЗАТЕЛЬНО экранируй:
   - Кавычки: \\"
   - Переводы строк: \\n
   - Обратный слеш: \\\\

Верни ТОЛЬКО валидный JSON, без markdown блоков и дополнительного текста."""

def format_duration(seconds: float) -> str:
    """Format duration in human-readable format"""
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = int(seconds // 60)
    secs = seconds % 60
    return f"{minutes}m {secs:.1f}s"


def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert DOCX to PDF using LibreOffice or python-docx + reportlab"""
    import tempfile
    import subprocess
    
    docx_path_obj = Path(docx_path)
    pdf_path = docx_path_obj.with_suffix('.pdf')

    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(docx_path_obj.parent), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=30
        )
        if pdf_path.exists():
            return str(pdf_path)
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT
        import re

        print(f"  Converting DOCX to PDF using python-docx + reportlab...")

        doc = Document(docx_path)

        font_registered = False
        font_name = None

        fonts_to_try = [
            ('/System/Library/Fonts/Supplemental/Arial Unicode.ttf', 'ArialUnicode'),
            ('/System/Library/Fonts/Helvetica.ttc', 'Helvetica'),
            ('/Library/Fonts/Arial.ttf', 'Arial'),
            ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 'DejaVuSans'),
            ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf', 'LiberationSans'),
            ('C:/Windows/Fonts/arial.ttf', 'Arial'),
            ('C:/Windows/Fonts/calibri.ttf', 'Calibri'),
        ]

        for font_path, fname in fonts_to_try:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(fname, font_path))
                    font_name = fname
                    font_registered = True
                    print(f"  ✓ Registered font: {fname}")
                    break
                except Exception as e:
                    continue

        if not font_registered:
            print("  ⚠ Warning: Could not register Unicode font, falling back to default (may not display Cyrillic properly)")
            font_name = 'Helvetica'

        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)

        elements = []
        styles = getSampleStyleSheet()

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10,  # Reduced from 11
            leading=12,   # Reduced from 14
            spaceAfter=3, # Reduced from 6
            alignment=TA_LEFT,
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=12,  # Reduced from 14
            leading=14,   # Reduced from 16
            spaceAfter=6, # Reduced from 12
            alignment=TA_LEFT,
        )

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text = text.replace('\n', '<br/>')
            text = text.replace('&', '&amp;')
            text = text.replace('<', '&lt;')
            text = text.replace('>', '&gt;')

            is_heading = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                    is_heading = True

            style = heading_style if is_heading else normal_style

            try:
                p = Paragraph(text, style)
                elements.append(p)
                if is_heading:
                    elements.append(Spacer(1, 3))
            except Exception as e:
                try:
                    p = Paragraph(text.replace('<br/>', ' '), normal_style)
                    elements.append(p)
                except:
                    print(f"  ⚠ Warning: Could not add paragraph: {text[:50]}...")
                    pass

        pdf_doc.build(elements)
        
        if Path(pdf_path).exists() and Path(pdf_path).stat().st_size > 0:
            print(f"  ✓ PDF created: {pdf_path.name}")
            return str(pdf_path)
        else:
            raise ValueError("PDF file was not created or is empty")
            
    except ImportError as e:
        print(f"  Missing dependencies: {e}")
        print("  Install: pip install python-docx reportlab")
        raise
    except Exception as e:
        print(f"  Error during conversion: {e}")
        raise
    
    raise ValueError(
        f"Cannot convert DOCX to PDF. Install dependencies:\n"
        f"  pip install python-docx reportlab\n"
        f"\nOr install LibreOffice (better quality):\n"
        f"  brew install --cask libreoffice"
    )


def merge_docx(docx_files: list[Path], output_path: str) -> None:
    """Merge multiple DOCX files into one"""
    from docx import Document

    if not docx_files:
        raise ValueError("No DOCX files to merge")

    merged_doc = Document()

    for i, docx_path in enumerate(docx_files):
        doc = Document(str(docx_path))

        if i > 0:
            merged_doc.add_page_break()

        from docx.oxml import parse_xml
        from docx.oxml.ns import qn

        for para in doc.paragraphs:
            para_xml_str = para._element.xml
            new_para_xml = parse_xml(para_xml_str)

            hyperlinks = new_para_xml.xpath('.//w:hyperlink')
            for hyperlink_elem in hyperlinks:
                rel_id = hyperlink_elem.get(qn('r:id'))
                if rel_id and rel_id in doc.part.rels:
                    source_rel = doc.part.rels[rel_id]
                    existing_rel_id = None
                    for r_id, r in merged_doc.part.rels.items():
                        if r.target_ref == source_rel.target_ref:
                            existing_rel_id = r_id
                            break

                    if existing_rel_id:
                        hyperlink_elem.set(qn('r:id'), existing_rel_id)
                    else:
                        new_rel = merged_doc.part.rels.get_or_add_ext_rel(
                            source_rel.reltype,
                            source_rel.target_ref
                        )
                        hyperlink_elem.set(qn('r:id'), new_rel.rId)

            merged_doc.element.body.append(new_para_xml)

        for table in doc.tables:
            new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
    
    merged_doc.save(output_path)


def merge_pdfs(pdf_files: list[Path], output_path: str) -> None:
    """Merge multiple PDF files into one"""
    try:
        from pypdf import PdfReader, PdfWriter
        
        writer = PdfWriter()
        for pdf_path in pdf_files:
            reader = PdfReader(str(pdf_path))
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
    except ImportError:
        try:
            from PyPDF2 import PdfMerger
            merger = PdfMerger()
            for pdf_path in pdf_files:
                merger.append(str(pdf_path))
            merger.write(output_path)
            merger.close()
        except ImportError:
            raise ImportError("Need pypdf or PyPDF2: pip install pypdf")


def merge_folder_to_docx(folder_path: str) -> str:
    """Merge all DOCX files from folder into one"""
    folder = Path(folder_path)

    docx_files = sorted([f for f in folder.glob('*.docx') if not f.name.startswith('~$')])

    if not docx_files:
        raise ValueError(f"No DOCX files found in {folder_path}")

    print(f"Found {len(docx_files)} DOCX files")

    merged_docx_name = f"{folder.name}_merged.docx"
    merged_docx_path = folder / merged_docx_name
    
    print(f"Merging {len(docx_files)} DOCX files into one...")
    merge_docx(docx_files, str(merged_docx_path))
    
    print(f"✓ Merged DOCX created: {merged_docx_path}")
    return str(merged_docx_path)


def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file"""
    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return "\n\n".join(text)
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(pdf_path)
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return "\n\n".join(text)
        except ImportError:
            raise ImportError("Need pypdf or PyPDF2: pip install pypdf")


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document

        doc = Document(docx_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text)
    except ImportError:
        raise ImportError("Need python-docx: pip install python-docx")


def extract_text(file_path: str) -> str:
    """Extract text from PDF or DOCX file"""
    file_path_obj = Path(file_path)

    if file_path_obj.suffix.lower() == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_path_obj.suffix.lower() == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path_obj.suffix}")


def upload_file_to_gemini(file_path: str):
    """Upload file (PDF or DOCX) to Gemini"""
    file_path_obj = Path(file_path)

    if file_path_obj.suffix.lower() == '.pdf':
        mime_type = 'application/pdf'
        display_name = 'quiz_pack.pdf'
        file_type = 'PDF'
    elif file_path_obj.suffix.lower() == '.docx':
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        display_name = 'quiz_pack.docx'
        file_type = 'DOCX'
    else:
        raise ValueError(f"Unsupported file type: {file_path_obj.suffix}. Supported: PDF, DOCX")

    print(f"Uploading {file_type}: {file_path_obj.name}")

    with open(file_path, 'rb') as f:
        file = client.files.upload(
            file=f,
            config=types.UploadFileConfig(
                mime_type=mime_type,
                display_name=display_name
            )
        )

    print(f"File uploaded: {file.name}")

    while file.state == types.FileState.PROCESSING:
        print("Processing file...")
        time.sleep(2)
        file = client.files.get(name=file.name)

    if file.state == types.FileState.FAILED:
        raise ValueError("File processing failed")

    print("File ready for analysis")
    return file

def extract_json_from_response(response_text: str):
    """Extract and parse JSON from Gemini response (returns dict or list)"""
    if response_text is None:
        raise ValueError("Response text is None - API may have failed or text was too large")

    response_text = response_text.strip()

    if response_text.startswith("```json"):
        response_text = response_text[7:]
    elif response_text.startswith("```"):
        response_text = response_text[3:]

    if response_text.endswith("```"):
        response_text = response_text[:-3]

    response_text = response_text.strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        print(f"Response text (first 1000 chars): {response_text[:1000]}")
        print(f"Response text (last 500 chars): {response_text[-500:]}")

        print("Attempting to fix JSON...")

        try:
            return json.loads(response_text, strict=False)
        except:
            pass

        debug_path = Path("debug_response.txt")
        debug_path.write_text(response_text, encoding='utf-8')
        print(f"Full response saved to: {debug_path}")

        raise ValueError(f"Failed to parse JSON. Saved full response to {debug_path}") from e

def get_package_structure(uploaded_file=None, text_content: str = None, is_merged: bool = False) -> dict:
    """Get package metadata and theme names (first pass)"""
    print("Step 1/2: Getting package structure...")

    merged_note = """
ОСОБЕННОСТИ ОБЪЕДИНЁННОГО ФАЙЛА (несколько "боёв"):
- Файл может содержать несколько "боёв" или "раундов" (разделены заголовками типа "Бой 1", "Бой 2")
- Парси ВСЕ темы из ВСЕХ боёв, не пропускай ничего
- Игнорируй колонтитулы, номера страниц, служебную информацию между боями
""" if is_merged else ""

    prompt = f"""Извлеки из {"текста" if text_content else "файла (PDF или DOCX)"} ТОЛЬКО метаданные и список тем (БЕЗ вопросов).

{merged_note}

Верни JSON:
{{
  "info": "полная информация о пакете: название, авторы, редакторы, благодарности, мораторий",
  "package_name": "название пакета",
  "themes": [
    {{"name": "Название темы 1 (Автор: Имя)"}},
    {{"name": "Название темы 2 (Автор: Имя)"}}
  ]
}}

ВАЖНО:
- НЕ включай вопросы, только названия тем
- Если информация о пакете (info) НЕ найдена в документе, оставь пустой строкой: ""
- Если название пакета (package_name) НЕ найдено, попробуй определить из содержимого (название файла, заголовки) или оставь пустой строкой: ""
- НЕ выдумывай информацию, которой нет в документе
- В JSON строках ОБЯЗАТЕЛЬНО экранируй кавычки как \\" и переводы строк как \\n
- Верни ТОЛЬКО валидный JSON без синтаксических ошибок"""

    # Build contents based on mode (file or text)
    if text_content:
        contents = [f"{prompt}\n\n--- ТЕКСТ ДОКУМЕНТА ---\n{text_content}"]
    else:
        contents = [
            types.Part.from_uri(file_uri=uploaded_file.uri, mime_type=uploaded_file.mime_type),
            prompt
        ]

    response = client.models.generate_content(
        model="gemini-2.5-pro",
        contents=contents,
        config=types.GenerateContentConfig(
            temperature=0.1,
            max_output_tokens=8192
        )
    )

    if not response.text:
        print("\n⚠️  ERROR: Response text is None!")
        print(f"Response object: {response}")

        if hasattr(response, 'candidates') and response.candidates:
            print(f"Number of candidates: {len(response.candidates)}")
            for i, candidate in enumerate(response.candidates):
                print(f"  Candidate {i}:")
                if hasattr(candidate, 'finish_reason'):
                    print(f"    Finish reason: {candidate.finish_reason}")
                if hasattr(candidate, 'safety_ratings'):
                    print(f"    Safety ratings: {candidate.safety_ratings}")
                if hasattr(candidate, 'content'):
                    print(f"    Has content: {candidate.content is not None}")

        if hasattr(response, 'prompt_feedback'):
            print(f"Prompt feedback: {response.prompt_feedback}")

        raise ValueError(
            "Response text is None. Possible causes:\n"
            "1. Content was blocked by safety filters\n"
            "2. Output exceeded max_output_tokens limit\n"
            "3. API error occurred\n"
            "Check the diagnostics above for details."
        )

    return extract_json_from_response(response.text)

def parse_theme_questions(uploaded_file=None, text_content: str = None, theme_name: str = None, theme_index: int = 0, is_merged: bool = False) -> list:
    """Parse questions for a specific theme (second pass)"""
    print(f"  Parsing theme {theme_index + 1}: {theme_name[:50]}...")

    merged_note = """
ОСОБЕННОСТИ ОБЪЕДИНЁННОГО ФАЙЛА:
- Если номера вопросов (10, 20, 30) не явные, определяй cost по порядку:
  * Первый вопрос в теме = cost: 10
  * Второй вопрос = cost: 20
  * И т.д.
""" if is_merged else ""

    prompt = f"""Извлеки данные для темы "{theme_name}" (тема #{theme_index + 1} в пакете).

ФОРМАТ ПОЛЯ "form" (маркер из вопроса, НЕ ответ!):
- "ОН был профессором" → form: "он"
- "В ЧЕСТЬ НЕГО назвали" → form: "в честь него"
- "ЭТОТ ФРАНЦУЗ" → form: "француз"
- "В ответе одно слово. ОНА" → form: "одним словом, она"
- Если в вопросе НЕТ заглавных местоимений, извлекай form из контекста:
  * "Что мы заменили" → form: "что мы заменили"
  * "Скажыце па-ўкраінскі" → form: "па-ўкраінскі"
  * "Назовите" → form: "назовите"
  * "Укажите" → form: "укажите"
  * "В ответе одно слово" → form: "одним словом"

{merged_note}

Верни JSON объект:
{{
  "theme_comment": "комментарий к теме, если есть в документе (НЕ выдумывай!)",
  "questions": [
    {{"cost": 10, "question": "...", "form": "он/она/его/город", "answer": "...", "source": "источник"}},
    {{"cost": 20, "question": "...", "form": "...", "answer": "...", "source": "..."}},
    ...
  ]
}}

ВАЖНО:
- theme_comment: добавляй ТОЛЬКО если в документе есть пояснение к теме (под названием). НЕ выдумывай!
- Вопросы должны быть отсортированы по cost (10, 20, 30, 40, 50)
- ВСЕГДА ищи и добавляй поле "source" (источник ответа) если оно указано в документе
- Источник обычно находится после ответа в скобках или отдельной строкой (примеры: "(источник: ...)", "Источник: ...", "[...]")
- Включай "comment" и "accept" только если они есть в PDF
- В JSON строках ОБЯЗАТЕЛЬНО экранируй кавычки как \\" и переводы строк как \\n
- Верни ТОЛЬКО валидный JSON объект без синтаксических ошибок"""

    # Build contents based on mode (file or text)
    if text_content:
        contents = [f"{prompt}\n\n--- ТЕКСТ ДОКУМЕНТА ---\n{text_content}"]
    else:
        contents = [
            types.Part.from_uri(file_uri=uploaded_file.uri, mime_type=uploaded_file.mime_type),
            prompt
        ]

    response = client.models.generate_content(
        model="gemini-2.5-pro",
        contents=contents,
        config=types.GenerateContentConfig(
            temperature=0.1,
            max_output_tokens=8192
        )
    )

    if not response.text:
        print(f"\n⚠️  ERROR: Response text is None for theme: {theme_name}")
        print(f"Response object: {response}")

        if hasattr(response, 'candidates') and response.candidates:
            print(f"Number of candidates: {len(response.candidates)}")
            for i, candidate in enumerate(response.candidates):
                print(f"  Candidate {i}:")
                if hasattr(candidate, 'finish_reason'):
                    print(f"    Finish reason: {candidate.finish_reason}")
                if hasattr(candidate, 'safety_ratings'):
                    print(f"    Safety ratings: {candidate.safety_ratings}")
                if hasattr(candidate, 'content'):
                    print(f"    Has content: {candidate.content is not None}")

        if hasattr(response, 'prompt_feedback'):
            print(f"Prompt feedback: {response.prompt_feedback}")

        raise ValueError(
            f"Response text is None for theme '{theme_name}'. Possible causes:\n"
            "1. Content was blocked by safety filters\n"
            "2. Output exceeded max_output_tokens limit\n"
            "3. API error occurred\n"
            "Check the diagnostics above for details."
        )

    result = extract_json_from_response(response.text)

    # Handle both old format (list) and new format (dict with theme_comment and questions)
    if isinstance(result, list):
        return {"questions": result}
    elif isinstance(result, dict) and "questions" in result:
        return result
    else:
        raise ValueError(f"Expected dict with questions or list, got {type(result)}")

def parse_with_text_chunked(file_path: str, is_merged: bool = False) -> dict:
    """Parse file using plain text extraction and Gemini API (faster, no file upload)"""

    total_start = time.time()

    extract_start = time.time()
    print(f"Extracting text from: {Path(file_path).name}")
    text_content = extract_text(file_path)
    extract_duration = time.time() - extract_start
    print(f"  ⏱ Text extraction took: {format_duration(extract_duration)}")
    print(f"  Text length: {len(text_content)} chars")

    structure_start = time.time()
    structure = get_package_structure(text_content=text_content, is_merged=is_merged)
    structure_duration = time.time() - structure_start
    print(f"  ⏱ Structure extraction took: {format_duration(structure_duration)}")

    if "themes" not in structure or not structure["themes"]:
        raise ValueError("Failed to extract themes from text")

    print(f"Found {len(structure['themes'])} themes")
    print(f"Step 2/2: Parsing questions for each theme...")

    themes_start = time.time()
    theme_times = []
    failed_themes = []

    for i, theme in enumerate(structure["themes"]):
        theme_start = time.time()
        theme_name = theme.get("name", f"Theme {i+1}")
        try:
            result = parse_theme_questions(text_content=text_content, theme_name=theme_name, theme_index=i, is_merged=is_merged)
            theme["questions"] = result.get("questions", [])
            if result.get("theme_comment"):
                theme["theme_comment"] = result["theme_comment"]
            theme_duration = time.time() - theme_start
            theme_times.append(theme_duration)
            print(f"    ✓ {len(theme['questions'])} questions parsed ({format_duration(theme_duration)})")
        except Exception as e:
            theme["questions"] = []
            failed_themes.append((i + 1, theme_name, str(e)))
            print(f"    ⚠️ FAILED: {theme_name[:50]}... - {e}")
            continue

    themes_duration = time.time() - themes_start
    total_duration = time.time() - total_start

    print()
    print("=" * 50)
    print("⏱ TIMING SUMMARY (TEXT MODE):")
    print(f"  Text extraction:     {format_duration(extract_duration)}")
    print(f"  Structure extraction: {format_duration(structure_duration)}")
    print(f"  Themes parsing:      {format_duration(themes_duration)}")
    if theme_times:
        avg_theme = sum(theme_times) / len(theme_times)
        print(f"    Avg per theme:     {format_duration(avg_theme)}")
    print(f"  TOTAL:               {format_duration(total_duration)}")
    if failed_themes:
        print(f"  ⚠️ FAILED THEMES:    {len(failed_themes)}")
        for num, name, err in failed_themes:
            print(f"    - Theme {num}: {name[:40]}...")
    print("=" * 50)
    print()

    return structure


def parse_pdf_with_gemini_chunked(file_path: str, is_merged: bool = False) -> dict:
    """Parse PDF/DOCX file using Gemini API with chunked approach (recommended for large files)"""

    total_start = time.time()

    upload_start = time.time()
    uploaded_file = upload_file_to_gemini(file_path)
    upload_duration = time.time() - upload_start
    print(f"  ⏱ Upload took: {format_duration(upload_duration)}")

    structure_start = time.time()
    structure = get_package_structure(uploaded_file, is_merged=is_merged)
    structure_duration = time.time() - structure_start
    print(f"  ⏱ Structure extraction took: {format_duration(structure_duration)}")

    if "themes" not in structure or not structure["themes"]:
        raise ValueError("Failed to extract themes from PDF")

    print(f"Found {len(structure['themes'])} themes")
    print(f"Step 2/2: Parsing questions for each theme...")

    themes_start = time.time()
    theme_times = []
    failed_themes = []

    for i, theme in enumerate(structure["themes"]):
        theme_start = time.time()
        theme_name = theme.get("name", f"Theme {i+1}")
        try:
            result = parse_theme_questions(uploaded_file, theme_name=theme_name, theme_index=i, is_merged=is_merged)
            theme["questions"] = result.get("questions", [])
            if result.get("theme_comment"):
                theme["theme_comment"] = result["theme_comment"]
            theme_duration = time.time() - theme_start
            theme_times.append(theme_duration)
            print(f"    ✓ {len(theme['questions'])} questions parsed ({format_duration(theme_duration)})")
        except Exception as e:
            theme["questions"] = []
            failed_themes.append((i + 1, theme_name, str(e)))
            print(f"    ⚠️ FAILED: {theme_name[:50]}... - {e}")
            continue

    themes_duration = time.time() - themes_start
    total_duration = time.time() - total_start

    print()
    print("=" * 50)
    print("⏱ TIMING SUMMARY (FILE MODE):")
    print(f"  Upload:              {format_duration(upload_duration)}")
    print(f"  Structure extraction: {format_duration(structure_duration)}")
    print(f"  Themes parsing:      {format_duration(themes_duration)}")
    if theme_times:
        avg_theme = sum(theme_times) / len(theme_times)
        print(f"    Avg per theme:     {format_duration(avg_theme)}")
    print(f"  TOTAL:               {format_duration(total_duration)}")
    if failed_themes:
        print(f"  ⚠️ FAILED THEMES:    {len(failed_themes)}")
        for num, name, err in failed_themes:
            print(f"    - Theme {num}: {name[:40]}...")
    print("=" * 50)
    print()

    return structure

def parse_pdf_with_gemini(file_path: str, chunked: bool = True, text_mode: bool = False, is_merged: bool = False) -> dict:
    """Parse PDF/DOCX file using Gemini API

    Args:
        file_path: Path to PDF or DOCX file
        chunked: If True, use chunked parsing (recommended for large files).
                 If False, use single-shot parsing (faster but may hit token limits).
        text_mode: If True, extract text and send as plain text (faster, no file upload).
                   If False, upload file to Gemini (supports images/formatting).
        is_merged: If True, file contains multiple "боёв" merged together
    """

    if text_mode:
        return parse_with_text_chunked(file_path, is_merged=is_merged)

    if chunked:
        return parse_pdf_with_gemini_chunked(file_path, is_merged=is_merged)

    uploaded_file = upload_file_to_gemini(file_path)

    print("Parsing file with Gemini (single-shot mode)...")

    prompt = """Распарси файл (PDF или DOCX) с вопросами "Своя игра" в JSON.

ФОРМАТ ПОЛЯ "form":
"form" - это маркер ЧТО назвать (ИЗ ВОПРОСА, НЕ ИЗ ОТВЕТА!):
- "ОН был профессором" → form: "он"
- "В ЧЕСТЬ НЕГО назвали" → form: "в честь него"
- "ЭТОТ ФРАНЦУЗ" → form: "француз"
- "В ответе одно слово. ОНА" → form: "одним словом, она"

JSON формат:
{"info":"...", "package_name":"...", "themes":[{"name":"...","questions":[{"cost":10,"question":"...","form":"он/она/его/город","answer":"...","source":"источник"}]}]}

ВАЖНО:
- ВСЕГДА ищи и добавляй поле "source" (источник ответа) если оно указано
- Опускай "comment" и "accept" если их нет
- Если информация о пакете (info) НЕ найдена в документе, оставь пустой строкой: ""
- Если название пакета (package_name) НЕ найдено, попробуй определить из содержимого или оставь пустой строкой: ""
- НЕ выдумывай информацию, которой нет в документе
- В JSON ОБЯЗАТЕЛЬНО экранируй кавычки (\\") и переводы строк (\\n)
- Будь краток. Закрой все скобки!"""

    response = client.models.generate_content(
        model="gemini-2.5-pro",
        contents=[
            types.Part.from_uri(file_uri=uploaded_file.uri, mime_type=uploaded_file.mime_type),
            prompt
        ],
        config=types.GenerateContentConfig(
            temperature=0.1,
            max_output_tokens=8192
        )
    )

    return extract_json_from_response(response.text)

def save_json(data: dict, output_path: str):
    """Save parsed data to JSON file"""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON saved to: {output_path}")

def validate_json_structure(data: dict) -> bool:
    """Validate that JSON has required structure and all questions have 'form' field"""
    if "info" not in data:
        print("Warning: Missing 'info' field")
        return False

    if "themes" not in data:
        print("Error: Missing 'themes' field")
        return False

    missing_forms = []
    empty_themes = []
    for theme_idx, theme in enumerate(data["themes"]):
        if "questions" not in theme:
            print(f"Error: Theme {theme_idx} missing 'questions' field")
            return False

        if len(theme["questions"]) == 0:
            empty_themes.append(f"Theme {theme_idx + 1}: {theme.get('name', 'Unknown')[:40]}...")
            continue

        for q_idx, question in enumerate(theme["questions"]):
            if "form" not in question:
                missing_forms.append(f"Theme {theme_idx}, Question {q_idx} (cost: {question.get('cost', '?')})")

    if empty_themes:
        print(f"⚠️ WARNING: {len(empty_themes)} theme(s) have no questions (failed to parse):")
        for item in empty_themes:
            print(f"  - {item}")

    if missing_forms:
        print("ERROR: Missing 'form' field in the following questions:")
        for item in missing_forms:
            print(f"  - {item}")
        return False

    total_questions = sum(len(t['questions']) for t in data['themes'])
    print("✓ JSON structure validation passed")
    print(f"✓ All {total_questions} questions have 'form' field")
    if empty_themes:
        print(f"⚠️ Note: {len(empty_themes)} theme(s) were skipped due to parsing errors")
    return True

def generate_short_name(text: str) -> str:
    """Generate slug from folder/package name

    Examples: "Дровушки_2022" -> "дровушки2022"
    """
    import re
    import unicodedata

    text = unicodedata.normalize('NFKD', text)

    if '.' in text:
        parts = text.split('.', 1)
        main_part = parts[0].strip()
        subtitle = parts[1].strip() if len(parts) > 1 else ''

        main_part = re.sub(r'[^\w\s]', '', main_part)
        main_part = main_part.lower().strip()
        main_part = re.sub(r'[\s_]+', '', main_part)

        if subtitle:
            subtitle_words = re.sub(r'[^\w\s]', '', subtitle).split()
            subtitle_initials = ''.join(word[0].lower() for word in subtitle_words if word)
            return main_part + subtitle_initials
        return main_part
    else:
        text = re.sub(r'[^\w\s]', '', text)
        text = text.lower().strip()
        text = re.sub(r'[\s_]+', '', text)
        return text

def main():
    import argparse

    parser = argparse.ArgumentParser(
        description='Parse quiz pack PDF/DOCX files using Gemini API',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Single PDF file
  python parse_pdf_with_gemini.py pack.pdf

  # Folder with multiple files (auto-detected)
  python parse_pdf_with_gemini.py Лагерь_Блик_2024_ЭК/

  # Explicit folder mode
  python parse_pdf_with_gemini.py Лагерь_Блик_2024_ЭК/ --folder

  # Use plain text mode (faster)
  python parse_pdf_with_gemini.py pack.docx --text-mode
        """
    )
    parser.add_argument('input_path', help='PDF file, DOCX file, or folder path')
    parser.add_argument('--folder', action='store_true',
                       help='Treat input_path as folder and merge files')
    parser.add_argument('--single-shot', action='store_true',
                       help='Use single-shot parsing (faster, but may hit limits)')
    parser.add_argument('--text-mode', action='store_true',
                       help='Extract plain text and send to Gemini (faster, no file upload)')

    args = parser.parse_args()

    input_path = Path(args.input_path)
    chunked = not args.single_shot
    text_mode = args.text_mode
    is_merged = False
    merged_pdf_path = None
    file_path = None

    if args.folder or (input_path.exists() and input_path.is_dir()):
        print(f"Folder mode: processing {input_path}")
        print("-" * 60)

        pdf_files = sorted([f for f in input_path.glob('*.pdf')])
        docx_files = sorted([f for f in input_path.glob('*.docx') if not f.name.startswith('~$')])

        if not pdf_files and not docx_files:
            print(f"❌ Error: No PDF or DOCX files found in {input_path}")
            sys.exit(1)

        print(f"Found {len(pdf_files)} PDF file(s) and {len(docx_files)} DOCX file(s)")

        all_pdfs = []
        temp_converted_pdfs = []

        if pdf_files:
            all_pdfs.extend(pdf_files)
            print(f"  ✓ {len(pdf_files)} PDF file(s) ready")

        if docx_files:
            print(f"  Converting {len(docx_files)} DOCX file(s) to PDF...")
            for i, docx_file in enumerate(docx_files, 1):
                print(f"    [{i}/{len(docx_files)}] Converting {docx_file.name}...")
                try:
                    converted_pdf = convert_docx_to_pdf(str(docx_file))
                    converted_pdf_path = Path(converted_pdf)
                    all_pdfs.append(converted_pdf_path)
                    temp_converted_pdfs.append(converted_pdf_path)
                    print(f"      ✓ Created: {converted_pdf_path.name}")
                except Exception as e:
                    print(f"      ⚠️  Failed to convert {docx_file.name}: {e}")
                    print(f"      Skipping this file...")

        if not all_pdfs:
            print(f"❌ Error: No valid PDF files to process")
            sys.exit(1)

        all_pdfs = sorted(all_pdfs, key=lambda p: p.name)

        merged_pdf_name = f"{input_path.name}_merged.pdf"
        merged_pdf_path = input_path / merged_pdf_name

        print(f"\nMerging {len(all_pdfs)} PDF file(s) into one...")
        merge_pdfs(all_pdfs, str(merged_pdf_path))
        print(f"✓ Merged PDF created: {merged_pdf_path}")

        file_path = str(merged_pdf_path)
        is_merged = True
    elif input_path.exists() and input_path.is_file():
        file_path_str = str(input_path)

        if input_path.suffix.lower() == '.docx' and not text_mode:
            print("Converting DOCX to PDF...")
            file_path = convert_docx_to_pdf(file_path_str)
        else:
            file_path = file_path_str
    else:
        print(f"Error: Path not found: {input_path}")
        sys.exit(1)

    if is_merged:
        output_path = input_path / f"{input_path.name}_parsed.json"
    else:
        output_path = Path(file_path).parent / f"{Path(file_path).stem}_parsed.json"
    
    print(f"Input: {input_path}")
    print(f"Output JSON: {output_path}")
    print(f"Parsing mode: {'Chunked (theme-by-theme)' if chunked else 'Single-shot'}")
    print(f"Processing mode: {'Plain text' if text_mode else 'File upload'}")
    print(f"File format: {Path(file_path).suffix.upper()}")
    if is_merged:
        print("Note: Processing merged file (multiple боёв)")
    print("-" * 60)

    try:
        result = parse_pdf_with_gemini(file_path, chunked=chunked, text_mode=text_mode, is_merged=is_merged)

        save_json(result, str(output_path))
        print()

        validation_passed = validate_json_structure(result)
        
        if not validation_passed:
            print("\n⚠️  WARNING: JSON validation failed!")
            print("The JSON structure is incomplete or missing 'form' fields.")
            print(f"⚠️  JSON saved anyway to: {output_path}")
            print("You can fix it manually or re-run the parser.")
            sys.exit(1)

        print("-" * 60)
        print("✓ Success! File parsed and saved to JSON")
        print(f"  Themes: {len(result['themes'])}")
        print(f"  Total questions: {sum(len(t['questions']) for t in result['themes'])}")
        if is_merged:
            print(f"  Merged DOCX saved: {file_path}")

        print()
        print("=" * 60)

        short_name = generate_short_name(input_path.name)

        package_name = result.get('package_name', '')

        command = f"python3 scripts/append_pack.py {short_name} {output_path}"
        if package_name:
            command += f' --name "{package_name}"'

        print("📋 Next step: Copy and run this command to add pack to database:")
        print()
        print(f"  {command}")
        print()
        print("=" * 60)

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
