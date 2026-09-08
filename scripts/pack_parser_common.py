"""Shared utilities for quiz pack parsing scripts."""

import json
import os
import re
import subprocess
import unicodedata
from pathlib import Path

SYSTEM_PROMPT = """Ты эксперт по парсингу PDF файлов с вопросами для игры "Своя игра" в формат JSON.

═══════════════════════════════════════════════════════════════
КРИТИЧЕСКИ ВАЖНО! ЛОГИКА ПОЛЯ "FORM":
═══════════════════════════════════════════════════════════════

⚠️ ПОЛЕ "form" ≠ ОТВЕТ! ⚠️

"form" - это маркер/указание ЧТО именно нужно назвать, извлекается ИЗ САМОГО ВОПРОСА.

═══════════════════════════════════════════════════════════════
АЛГОРИТМ ИЗВЛЕЧЕНИЯ "FORM":
═══════════════════════════════════════════════════════════════

1️⃣ ИЩЕМ ЗАГЛАВНЫЕ СЛОВА/МЕСТОИМЕНИЯ В ВОПРОСЕ:
   - "В ЧЕСТЬ НЕГО назвали ящерицу" → form: "в честь него"
   - "ЭТОГО ГОЛЛАНДЦА упоминают" → form: "голландец" / "этого голландца"
   - "ОН шутил, что хотел" → form: "он"
   - "ЕГО хотели создать" → form: "его"
   - "ОНИ подвергались гонениям" → form: "они"
   - "С НЕЙ связано много легенд" → form: "она"
   - "ИМИ оказываются девушки" → form: "ими"
   - "ТАКИЕ ОНИ используются в игре" → form: "такие они"
   - "ЭТОЙ СТОЛИЦЕЙ является" → form: "столица" / "эта столица"
   - "ЭТОТ ГОРОД стал побратимом" → form: "этот город" / "город"
   - "В ЭТОМ ГОДУ погибла Ида" → form: "год"
   - "ЭТОГО ПРОДЮСЕРА работой было" → form: "продюсер"

2️⃣ ДОБАВЛЯЕМ ОГРАНИЧЕНИЯ ИЗ НАЧАЛА ВОПРОСА:
   - "В ответе одно слово. ОНА называется" → form: "одним словом, она"
   - "В ответе два слова. ОН известен" → form: "два слова, он"
   - "В ответе три слова. ЕГО называют" → form: "тремя словами, его"
   - "В ответе трёхсложное слово. ОНО" → form: "трёхсложное слово, оно"
   - "В ответе два слова на одну букву" → form: "два слова на одну букву, им"
   - "В ответе два слова на парные согласные" → form: "два слова на парные согласные"
   - "Зачет абсолютно точный ответ" → добавь это в form

3️⃣ ИСПОЛЬЗУЕМ ТИП ОБЪЕКТА ИЗ ВОПРОСА (если указан явно):
   - "ЭТОТ ФРАНЦУЗ не обладал" → form: "француз"
   - "ЭТА СТОЛИЦА находится" → form: "столица"
   - "ЭТОТ ГОРОД США стал" → form: "этот город"
   - "ЭТИМ АНГЛИЙСКИМ ВЫРАЖЕНИЕМ называют" → form: "английское выражение"
   - "ЭТОМУ ПРОИЗВЕДЕНИЮ посвящено" → form: "произведение"

4️⃣ ЕСЛИ В PDF ЕСТЬ СТРОКА "Зачет:":
   - Извлеки альтернативные варианты в отдельное поле "accept": ["вариант1", "вариант2"]
   - НО form всё равно должна быть местоимением/маркером из вопроса!

═══════════════════════════════════════════════════════════════
РЕАЛЬНЫЕ ПРИМЕРЫ ИЗ ПАКЕТОВ:
═══════════════════════════════════════════════════════════════

✅ ПРАВИЛЬНО:
Вопрос: "В декабре 2012 года В ЧЕСТЬ НЕГО назвали ископаемую ящерицу"
→ form: "в честь него"
→ answer: "Барак Обама"

✅ ПРАВИЛЬНО:
Вопрос: "Эстонец упоминается наравне с фамилией ЭТОГО ГОЛЛАНДЦА"
→ form: "голландец"
→ answer: "Ян Хендрик Оорт"

✅ ПРАВИЛЬНО:
Вопрос: "В ответе одно слово. «Третьей ЕЙ» называют загрязнение"
→ form: "одним словом, она"
→ answer: "рука"

✅ ПРАВИЛЬНО:
Вопрос: "ОН был профессором кафедры биохимии"
→ form: "он"
→ answer: "Александр Опарин"

✅ ПРАВИЛЬНО:
Вопрос: "ЭТА СТОЛИЦА находится в 25 км от экватора"
→ form: "столица"
→ answer: "Кито"

✅ ПРАВИЛЬНО:
Вопрос: "Девушки оказываются ИМИ"
→ form: "ими"
→ answer: "вампиры"

❌ НЕПРАВИЛЬНО:
→ form: "Барак Обама" ❌ (это ответ, не форма!)
→ form: "ответ" ❌ (бессмысленно)

═══════════════════════════════════════════════════════════════
ФОРМАТ ВЫХОДНОГО JSON:
═══════════════════════════════════════════════════════════════

{
  "info": "полная информация о пакете: название, авторы, редакторы, благодарности, мораторий",
  "package_name": "название пакета",
  "themes": [
    {
      "name": "Название темы (Автор: Имя Фамилия)",
      "theme_comment": "комментарий к теме, если есть в документе (НЕ выдумывать!)",
      "questions": [
        {
          "cost": 10,
          "question": "текст вопроса с ЗАГЛАВНЫМИ местоимениями",
          "form": "местоимение/маркер из вопроса (ОН/ОНА/ЕГО/ИХ/город/столица и т.д.)",
          "answer": "сам ответ",
          "source": "источник ответа (если указан)",
          "comment": "комментарий к ответу",
          "accept": ["альтернатива1", "альтернатива2"]
        }
      ]
    }
  ]
}

═══════════════════════════════════════════════════════════════
ОБЩИЕ ПРАВИЛА ПАРСИНГА:
═══════════════════════════════════════════════════════════════

1. Структура: Бои/Раунды → Темы → Вопросы (10, 20, 30, 40, 50 очков)
2. Info извлекается из начала файла (авторы, редакторы, благодарности, мораторий)
3. Сохраняй все ударения, специальные символы, форматирование
4. Поддерживай белорусский язык ("Адказ:", "Каментар:")
5. Если в вопросе несколько заглавных местоимений, выбирай основное (обычно последнее перед вопросительным знаком)
6. ВСЕГДА ищи и добавляй "source" (источник ответа) если он указан в документе:
   - Источник обычно после ответа в скобках: "(источник: ...)", "Источник: ...", "[...]"
   - Если источника нет, можно опустить поле
7. КРИТИЧНО: В JSON строках ОБЯЗАТЕЛЬНО экранируй:
   - Кавычки: \\"
   - Переводы строк: \\n
   - Обратный слеш: \\\\

Верни ТОЛЬКО валидный JSON, без markdown блоков и дополнительного текста."""


def format_duration(seconds: float) -> str:
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = int(seconds // 60)
    secs = seconds % 60
    return f"{minutes}m {secs:.1f}s"


def merged_structure_note(is_merged: bool) -> str:
    if not is_merged:
        return ""
    return """
ОСОБЕННОСТИ ОБЪЕДИНЁННОГО ФАЙЛА (несколько "боёв"):
- Файл может содержать несколько "боёв" или "раундов" (разделены заголовками типа "Бой 1", "Бой 2")
- Парси ВСЕ темы из ВСЕХ боёв, не пропускай ничего
- Игнорируй колонтитулы, номера страниц, служебную информацию между боями
"""


def merged_theme_note(is_merged: bool) -> str:
    if not is_merged:
        return ""
    return """
ОСОБЕННОСТИ ОБЪЕДИНЁННОГО ФАЙЛА:
- Если номера вопросов (10, 20, 30) не явные, определяй cost по порядку:
  * Первый вопрос в теме = cost: 10
  * Второй вопрос = cost: 20
  * И т.д.
"""


def build_structure_prompt(is_merged: bool = False) -> str:
    merged_note = merged_structure_note(is_merged)
    return f"""Извлеки из текста документа ТОЛЬКО метаданные и список тем (БЕЗ вопросов).

{merged_note}

Верни JSON:
{{
  "info": "полная информация о пакете: название, авторы, редакторы, благодарности, мораторий",
  "package_name": "название пакета",
  "themes": [
    {{"name": "Название темы 1 (Автор: Имя)"}},
    {{"name": "Название темы 2 (Автор: Имя)"}}
  ]
}}

ВАЖНО:
- НЕ включай вопросы, только названия тем
- Если информация о пакете (info) НЕ найдена в документе, оставь пустой строкой: ""
- Если название пакета (package_name) НЕ найдено, попробуй определить из содержимого (название файла, заголовки) или оставь пустой строкой: ""
- НЕ выдумывай информацию, которой нет в документе
- В JSON строках ОБЯЗАТЕЛЬНО экранируй кавычки как \\" и переводы строк как \\n
- Верни ТОЛЬКО валидный JSON без синтаксических ошибок"""


def build_theme_prompt(theme_name: str, theme_index: int, is_merged: bool = False) -> str:
    merged_note = merged_theme_note(is_merged)
    return f"""Извлеки данные для темы "{theme_name}" (тема #{theme_index + 1} в пакете).

ФОРМАТ ПОЛЯ "form" (маркер из вопроса, НЕ ответ!):
- "ОН был профессором" → form: "он"
- "В ЧЕСТЬ НЕГО назвали" → form: "в честь него"
- "ЭТОТ ФРАНЦУЗ" → form: "француз"
- "В ответе одно слово. ОНА" → form: "одним словом, она"
- Если в вопросе НЕТ заглавных местоимений, извлекай form из контекста:
  * "Что мы заменили" → form: "что мы заменили"
  * "Скажыце па-ўкраінскі" → form: "па-ўкраінскі"
  * "Назовите" → form: "назовите"
  * "Укажите" → form: "укажите"
  * "В ответе одно слово" → form: "одним словом"

{merged_note}

Верни JSON объект:
{{
  "theme_comment": "комментарий к теме, если есть в документе (НЕ выдумывай!)",
  "questions": [
    {{"cost": 10, "question": "...", "form": "он/она/его/город", "answer": "...", "source": "источник"}},
    {{"cost": 20, "question": "...", "form": "...", "answer": "...", "source": "..."}}
  ]
}}

ВАЖНО:
- theme_comment: добавляй ТОЛЬКО если в документе есть пояснение к теме (под названием). НЕ выдумывай!
- Вопросы должны быть отсортированы по cost (10, 20, 30, 40, 50)
- ВСЕГДА ищи и добавляй поле "source" (источник ответа) если оно указано в документе
- Источник обычно находится после ответа в скобках или отдельной строкой (примеры: "(источник: ...)", "Источник: ...", "[...]")
- Включай "comment" и "accept" только если они есть в PDF
- В JSON строках ОБЯЗАТЕЛЬНО экранируй кавычки как \\" и переводы строк как \\n
- Верни ТОЛЬКО валидный JSON объект без синтаксических ошибок"""


def build_user_message(prompt: str, text_content: str) -> str:
    return f"{prompt}\n\n--- ТЕКСТ ДОКУМЕНТА ---\n{text_content}"


def convert_docx_to_pdf(docx_path: str) -> str:
    docx_path_obj = Path(docx_path)
    pdf_path = docx_path_obj.with_suffix('.pdf')

    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(docx_path_obj.parent), str(docx_path)],
            check=True,
            capture_output=True,
            timeout=30
        )
        if pdf_path.exists():
            return str(pdf_path)
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        pass

    try:
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT

        print("  Converting DOCX to PDF using python-docx + reportlab...")

        doc = Document(docx_path)

        font_registered = False
        font_name = None

        fonts_to_try = [
            ('/System/Library/Fonts/Supplemental/Arial Unicode.ttf', 'ArialUnicode'),
            ('/System/Library/Fonts/Helvetica.ttc', 'Helvetica'),
            ('/Library/Fonts/Arial.ttf', 'Arial'),
            ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 'DejaVuSans'),
            ('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf', 'LiberationSans'),
            ('C:/Windows/Fonts/arial.ttf', 'Arial'),
            ('C:/Windows/Fonts/calibri.ttf', 'Calibri'),
        ]

        for font_path, fname in fonts_to_try:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(fname, font_path))
                    font_name = fname
                    font_registered = True
                    print(f"  ✓ Registered font: {fname}")
                    break
                except Exception:
                    continue

        if not font_registered:
            print("  ⚠ Warning: Could not register Unicode font, falling back to default (may not display Cyrillic properly)")
            font_name = 'Helvetica'

        pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)

        elements = []
        styles = getSampleStyleSheet()

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10,
            leading=12,
            spaceAfter=3,
            alignment=TA_LEFT,
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=font_name,
            fontSize=12,
            leading=14,
            spaceAfter=6,
            alignment=TA_LEFT,
        )

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text = text.replace('\n', '<br/>')
            text = text.replace('&', '&amp;')
            text = text.replace('<', '&lt;')
            text = text.replace('>', '&gt;')

            is_heading = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                    is_heading = True

            style = heading_style if is_heading else normal_style

            try:
                p = Paragraph(text, style)
                elements.append(p)
                if is_heading:
                    elements.append(Spacer(1, 3))
            except Exception:
                try:
                    p = Paragraph(text.replace('<br/>', ' '), normal_style)
                    elements.append(p)
                except Exception:
                    print(f"  ⚠ Warning: Could not add paragraph: {text[:50]}...")

        pdf_doc.build(elements)

        if Path(pdf_path).exists() and Path(pdf_path).stat().st_size > 0:
            print(f"  ✓ PDF created: {pdf_path.name}")
            return str(pdf_path)
        raise ValueError("PDF file was not created or is empty")

    except ImportError as e:
        print(f"  Missing dependencies: {e}")
        print("  Install: pip install python-docx reportlab")
        raise
    except Exception as e:
        print(f"  Error during conversion: {e}")
        raise

    raise ValueError(
        "Cannot convert DOCX to PDF. Install dependencies:\n"
        "  pip install python-docx reportlab\n"
        "\nOr install LibreOffice (better quality):\n"
        "  brew install --cask libreoffice"
    )


def merge_docx(docx_files: list[Path], output_path: str) -> None:
    from docx import Document

    if not docx_files:
        raise ValueError("No DOCX files to merge")

    merged_doc = Document()

    for i, docx_path in enumerate(docx_files):
        doc = Document(str(docx_path))

        if i > 0:
            merged_doc.add_page_break()

        from docx.oxml import parse_xml
        from docx.oxml.ns import qn

        for para in doc.paragraphs:
            para_xml_str = para._element.xml
            new_para_xml = parse_xml(para_xml_str)

            hyperlinks = new_para_xml.xpath('.//w:hyperlink')
            for hyperlink_elem in hyperlinks:
                rel_id = hyperlink_elem.get(qn('r:id'))
                if rel_id and rel_id in doc.part.rels:
                    source_rel = doc.part.rels[rel_id]
                    existing_rel_id = None
                    for r_id, r in merged_doc.part.rels.items():
                        if r.target_ref == source_rel.target_ref:
                            existing_rel_id = r_id
                            break

                    if existing_rel_id:
                        hyperlink_elem.set(qn('r:id'), existing_rel_id)
                    else:
                        new_rel = merged_doc.part.rels.get_or_add_ext_rel(
                            source_rel.reltype,
                            source_rel.target_ref
                        )
                        hyperlink_elem.set(qn('r:id'), new_rel.rId)

            merged_doc.element.body.append(new_para_xml)

        for table in doc.tables:
            new_table = merged_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            for row_i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[row_i].cells[j].text = cell.text

    merged_doc.save(output_path)


def merge_pdfs(pdf_files: list[Path], output_path: str) -> None:
    try:
        from pypdf import PdfReader, PdfWriter

        writer = PdfWriter()
        for pdf_path in pdf_files:
            reader = PdfReader(str(pdf_path))
            for page in reader.pages:
                writer.add_page(page)
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
    except ImportError:
        try:
            from PyPDF2 import PdfMerger
            merger = PdfMerger()
            for pdf_path in pdf_files:
                merger.append(str(pdf_path))
            merger.write(output_path)
            merger.close()
        except ImportError:
            raise ImportError("Need pypdf or PyPDF2: pip install pypdf")


def extract_text_from_pdf(pdf_path: str) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(pdf_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return "\n\n".join(text)
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(pdf_path)
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            return "\n\n".join(text)
        except ImportError:
            raise ImportError("Need pypdf or PyPDF2: pip install pypdf")


def extract_text_from_docx(docx_path: str) -> str:
    try:
        from docx import Document

        doc = Document(docx_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        return "\n".join(text)
    except ImportError:
        raise ImportError("Need python-docx: pip install python-docx")


def extract_text(file_path: str) -> str:
    file_path_obj = Path(file_path)

    if file_path_obj.suffix.lower() == '.pdf':
        return extract_text_from_pdf(file_path)
    if file_path_obj.suffix.lower() == '.docx':
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {file_path_obj.suffix}")


def extract_json_from_response(response_text: str):
    if response_text is None:
        raise ValueError("Response text is None - API may have failed or text was too large")

    response_text = response_text.strip()

    if response_text.startswith("```json"):
        response_text = response_text[7:]
    elif response_text.startswith("```"):
        response_text = response_text[3:]

    if response_text.endswith("```"):
        response_text = response_text[:-3]

    response_text = response_text.strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        print(f"Response text (first 1000 chars): {response_text[:1000]}")
        print(f"Response text (last 500 chars): {response_text[-500:]}")
        print("Attempting to fix JSON...")

        try:
            return json.loads(response_text, strict=False)
        except Exception:
            pass

        debug_path = Path("debug_response.txt")
        debug_path.write_text(response_text, encoding='utf-8')
        print(f"Full response saved to: {debug_path}")

        raise ValueError(f"Failed to parse JSON. Saved full response to {debug_path}") from e


def save_json(data: dict, output_path: str):
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"JSON saved to: {output_path}")


def validate_json_structure(data: dict) -> bool:
    if "info" not in data:
        print("Warning: Missing 'info' field")
        return False

    if "themes" not in data:
        print("Error: Missing 'themes' field")
        return False

    missing_forms = []
    empty_themes = []
    for theme_idx, theme in enumerate(data["themes"]):
        if "questions" not in theme:
            print(f"Error: Theme {theme_idx} missing 'questions' field")
            return False

        if len(theme["questions"]) == 0:
            empty_themes.append(f"Theme {theme_idx + 1}: {theme.get('name', 'Unknown')[:40]}...")
            continue

        for q_idx, question in enumerate(theme["questions"]):
            if "form" not in question:
                missing_forms.append(f"Theme {theme_idx}, Question {q_idx} (cost: {question.get('cost', '?')})")

    if empty_themes:
        print(f"⚠️ WARNING: {len(empty_themes)} theme(s) have no questions (failed to parse):")
        for item in empty_themes:
            print(f"  - {item}")

    if missing_forms:
        print("ERROR: Missing 'form' field in the following questions:")
        for item in missing_forms:
            print(f"  - {item}")
        return False

    total_questions = sum(len(t['questions']) for t in data['themes'])
    print("✓ JSON structure validation passed")
    print(f"✓ All {total_questions} questions have 'form' field")
    if empty_themes:
        print(f"⚠️ Note: {len(empty_themes)} theme(s) were skipped due to parsing errors")
    return True


def generate_short_name(text: str) -> str:
    text = unicodedata.normalize('NFKD', text)

    if '.' in text:
        parts = text.split('.', 1)
        main_part = parts[0].strip()
        subtitle = parts[1].strip() if len(parts) > 1 else ''

        main_part = re.sub(r'[^\w\s]', '', main_part)
        main_part = main_part.lower().strip()
        main_part = re.sub(r'[\s_]+', '', main_part)

        if subtitle:
            subtitle_words = re.sub(r'[^\w\s]', '', subtitle).split()
            subtitle_initials = ''.join(word[0].lower() for word in subtitle_words if word)
            return main_part + subtitle_initials
        return main_part

    text = re.sub(r'[^\w\s]', '', text)
    text = text.lower().strip()
    text = re.sub(r'[\s_]+', '', text)
    return text
